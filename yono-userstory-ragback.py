import os
import io
import re
import json
import time
import uuid
import base64
import tempfile
from typing import TypedDict, List, Dict, Any, Optional, Tuple

import httpx
import pandas as pd
import streamlit as st
import pdfplumber
from docx import Document
from PIL import Image

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter


BASE_URL = "https://genailab.tcs.in"
API_KEY = "sk-ClT4kUp-h82oJ-TZ9DYTQw"

tiktoken_cache_dir = "./token"
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir 

client = httpx.Client(verify=False) 


st.set_page_config(page_title="Compliance Reviewer AI", layout="wide")
st.title("📘 Back Office Compliance Document Reviewer")
st.caption("Upload compliance documents, detect risks, extract key clauses, and generate review highlights.")


planner_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure/genailab-maas-gpt-4o",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

reasoning_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure_ai/genailab-maas-DeepSeek-V3-0324",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

critic_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure_ai/genailab-maas-DeepSeek-V3-0324",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

final_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure/genailab-maas-gpt-4o-mini",
    api_key=API_KEY,
    http_client=client,
    temperature=0.2
)

vision_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure/genailab-maas-gpt-4o",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

embedding_model = OpenAIEmbeddings(
    base_url=BASE_URL,
    model="azure/genailab-maas-text-embedding-3-large",
    api_key=API_KEY,
    http_client=client
)


class AgentState(TypedDict):
    user_query: str
    document_name: str
    document_type: str
    plan: Dict[str, Any]
    active_rules: List[Dict[str, Any]]
    retrieved_chunks: List[str]
    retrieved_metadata: List[Dict[str, Any]]
    vision_results: List[str]
    rule_findings: List[Dict[str, Any]]
    reasoning_output: str
    critic_feedback: Dict[str, Any]
    final_answer: str
    csv_summary: Optional[Dict[str, Any]]
    


def safe_json_loads(text: str, fallback: Optional[dict] = None) -> dict:
    try:
        cleaned = text.strip()

        if cleaned.startswith("```json"):
            cleaned = cleaned[len("```json"):].strip()
        elif cleaned.startswith("```"):
            cleaned = cleaned[len("```"):].strip()

        if cleaned.endswith("```"):
            cleaned = cleaned[:-3].strip()

        return json.loads(cleaned)
    except Exception:
        return fallback or {}


def split_text(text: str) -> List[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=250,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    return splitter.split_text(text)

def csv_precheck(df: pd.DataFrame) -> Dict[str, Any]:
    working_df = df.copy()

    missing_values = {}
    for col in working_df.columns:
        missing_count = working_df[col].isna().sum()
        blank_count = (working_df[col].astype(str).str.strip() == "").sum()
        missing_values[str(col)] = int(max(missing_count, blank_count))

    report = {
        "row_count": int(len(working_df)),
        "column_count": int(len(working_df.columns)),
        "columns": [str(c) for c in working_df.columns.tolist()],
        "missing_values": missing_values,
        "duplicate_rows": int(working_df.duplicated().sum())
    }
    return report
def evaluate_csv_structure_rules(
    df: pd.DataFrame,
    rules: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    findings = []
    columns_lower = [str(c).lower() for c in df.columns]

    for rule in rules:
        required_columns = [str(c).lower() for c in rule.get("required_columns", [])]
        if not required_columns:
            continue

        missing_cols = [c for c in required_columns if c not in columns_lower]

        findings.append({
            "rule_id": rule.get("id", ""),
            "rule_name": rule.get("name", ""),
            "matched": len(missing_cols) > 0,
            "severity": rule.get("severity", "Medium"),
            "finding": (
                f"Missing required columns: {missing_cols}"
                if missing_cols else
                "All required columns are present."
            ),
            "evidence_summary": f"CSV columns: {df.columns.tolist()}",
            "risk_reason": rule.get("risk_if_missing", ""),
            "reviewer_recommendation": rule.get("reviewer_recommendation", "")
        })

    return findings
def create_vector_db(chunks: List[str], metadatas: Optional[List[dict]] = None):
    persist_dir = os.path.join(tempfile.gettempdir(), f"compliance_chroma_{uuid.uuid4().hex}")
    return Chroma.from_texts(
        texts=chunks,
        embedding=embedding_model,
        metadatas=metadatas or [{} for _ in chunks],
        persist_directory=persist_dir
    )


def encode_pil_image(image: Image.Image) -> str:
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()



def extract_image_with_llm(image: Image.Image) -> str:
    img_b64 = encode_pil_image(image)

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Extract and explain all meaningful content from this compliance-related image. "
                        "Capture headings, clauses, tables, steps, regulatory references, risk-related statements, "
                        "and anything important for compliance review."
                    )
                },
                {
                    "type": "image_url",
                    "image_url": f"data:image/png;base64,{img_b64}"
                }
            ]
        }
    ]

    response = vision_llm.invoke(messages)
    return response.content

def extract_csv(file) -> Tuple[str, List[Image.Image], Optional[pd.DataFrame]]:
    file.seek(0)
    raw = file.read()

    for encoding in ["utf-8", "utf-8-sig", "latin1", "cp1252"]:
        try:
            decoded = raw.decode(encoding)
            break
        except Exception:
            continue
    else:
        decoded = raw.decode("utf-8", errors="ignore")

    df = pd.read_csv(io.StringIO(decoded), sep=None, engine="python")
    df = df.fillna("")

    content = []
    content.append("[CSV Columns]")
    content.append(", ".join(df.columns.astype(str).tolist()))
    content.append(f"[Total Rows] {len(df)}")

    for idx, row in df.iterrows():
        row_dict = {str(col): str(row[col]) for col in df.columns}
        content.append(f"[Row {idx + 1}] " + json.dumps(row_dict, ensure_ascii=False))

    extracted_text = "\n".join(content)
    return extracted_text, [], df

def build_csv_chunks(df: pd.DataFrame, source_name: str, max_rows: int = 200):
    chunks = []
    metadatas = []

    working_df = df.fillna("").head(max_rows)

    for idx, row in working_df.iterrows():
        row_dict = {str(col): str(row[col]) for col in df.columns}
        chunk_text = f"CSV Row {idx + 1}: " + json.dumps(row_dict, ensure_ascii=False)
        chunks.append(chunk_text)
        metadatas.append({
            "source": source_name,
            "chunk_id": idx + 1,
            "row_number": idx + 1,
            "content_type": "csv_row"
        })

    return chunks, metadatas

def extract_text_from_rules_pdf(pdf_file) -> str:
    """Extract raw text from the uploaded rules PDF using pdfplumber."""
    content = []
    with pdfplumber.open(pdf_file) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                content.append(f"[Page {page_num}]\n{text}")
    return "\n\n".join(content)


def deepseek_rules_extraction_agent(pdf_text: str) -> List[Dict[str, Any]]:
    """
    Uses DeepSeek (reasoning_llm) as an agent to parse raw PDF text
    and extract structured compliance rules as a JSON list.
    """

    prompt = f"""You are a compliance rules extraction agent powered by DeepSeek.

You are given raw text extracted from a compliance rules PDF document.
Your job is to extract ALL compliance rules from this text and return them
as a structured JSON array.

Each rule MUST follow this exact schema:
{{
  "id": "R001",                          // unique rule ID, e.g. R001, R002, ...
  "name": "Short rule name",
  "category": "Category name",
  "severity": "High | Medium | Low",
  "description": "What this rule checks for",
  "required_keywords": ["keyword1", "keyword2"],   // words that should be present
  "prohibited_patterns": ["pattern1"],             // phrases that indicate a violation (can be empty list)
  "risk_if_missing": "What risk exists if rule is violated",
  "reviewer_recommendation": "What the reviewer should do"
}}

Rules for extraction:
- Extract EVERY rule you can identify in the document
- If the PDF has numbered rules, sections, clauses, or policies — treat each as one rule
- Infer severity from language: "must", "mandatory", "critical" = High; "should", "recommended" = Medium; "may", "optional" = Low
- If a field cannot be determined, use a reasonable default (e.g. empty list for keywords)
- Return ONLY the JSON array. No markdown. No explanation. No preamble.

Raw PDF text:
{pdf_text}
"""

    response = reasoning_llm.invoke(prompt)
    raw = response.content.strip()

    # Strip markdown fences if present
    if raw.startswith("```json"):
        raw = raw[len("```json"):].strip()
    elif raw.startswith("```"):
        raw = raw[len("```"):].strip()
    if raw.endswith("```"):
        raw = raw[:-3].strip()

    try:
        rules = json.loads(raw)
        if isinstance(rules, list):
            return rules
        elif isinstance(rules, dict) and "rules" in rules:
            return rules["rules"]
        else:
            st.warning("DeepSeek returned unexpected JSON structure. Wrapped in list.")
            return [rules]
    except json.JSONDecodeError as e:
        st.error(f"DeepSeek rules extraction failed to parse JSON: {e}")
        return []



def extract_docx(file) -> Tuple[str, List[Image.Image]]:
    doc = Document(file)
    content = []
    images = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            content.append(txt)

    for table in doc.tables:
        content.append("[Table]")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                content.append(" | ".join(row_text))

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target.blob
                image = Image.open(io.BytesIO(image_data)).convert("RGB")
                images.append(image)

                image_text = extract_image_with_llm(image)
                if image_text:
                    content.append("[Image Understanding]")
                    content.append(image_text)
            except Exception:
                continue

    return "\n".join(content), images, None


def extract_pdf(file) -> Tuple[str, List[Image.Image]]:
    content = []
    images = []

    with pdfplumber.open(file) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                content.append(f"[PDF Page {page_num}]")
                content.append(text)

    return "\n".join(content), images, None


def extract_txt(file) -> Tuple[str, List[Image.Image]]:
    text = file.read().decode("utf-8", errors="ignore")
    return text, [], None


def extract_content(uploaded_file) -> Tuple[str, List[Image.Image]]:
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_pdf(uploaded_file)
    if name.endswith(".docx"):
        return extract_docx(uploaded_file)
    if name.endswith(".txt"):
        return extract_txt(uploaded_file)
    if name.endswith(".csv"):
        return extract_csv(uploaded_file)

    raise ValueError("Unsupported file type. Use PDF, CSV, DOCX, or TXT.")



def format_retrieved_chunks(chunks: List[str], metadata: List[Dict[str, Any]]) -> str:
    formatted = []
    for i, chunk in enumerate(chunks):
        meta = metadata[i] if i < len(metadata) else {}
        source = meta.get("source", "unknown")
        chunk_id = meta.get("chunk_id", i + 1)
        formatted.append(f"[Chunk {i+1} | Source: {source} | Chunk ID: {chunk_id}]\n{chunk}")
    return "\n\n".join(formatted)


def planner_agent(state: AgentState) -> AgentState:

    prompt = f"""
        You are a planner agent for a compliance document review system.

        User request:
        {state['user_query']}

        Return ONLY valid JSON. No markdown. No explanation.

        JSON format:
        {{
        "needs_retrieval": true,
        "needs_vision": false,
        "task_breakdown": [
            "real task 1",
            "real task 2",
            "..."
        ],
        "review_focus": [
            "risk detection",
            "key clauses",
            "missing compliance language",
            "regulatory ambiguity"
        ],
        "answer_style": "concise"
        }}

        Rules:
        - task_breakdown can contain any number of real tasks
        - set needs_vision=true only if the query refers to image, figure, screenshot, scan, or diagram
        - focus on compliance review and risk analysis
        """

    response = planner_llm.invoke(prompt)

    state["plan"] = safe_json_loads(
        response.content,
        fallback={
            "needs_retrieval": True,
            "needs_vision": False,
            "task_breakdown": [
                "summarize the compliance document",
                "identify compliance risks",
                "highlight key clauses"
            ],
            "review_focus": [
                "risk detection",
                "key clauses",
                "ambiguity"
            ],
            "answer_style": "simple and easy to read"
        }
    )

    return state


def retriever_agent(state: AgentState, vectordb: Chroma, k: int = 5) -> AgentState:

    retriever = vectordb.as_retriever(search_kwargs={"k": k})
    docs = retriever.invoke(state["user_query"])

    state["retrieved_chunks"] = [doc.page_content for doc in docs]
    state["retrieved_metadata"] = [doc.metadata for doc in docs]

    return state


def vision_agent(state: AgentState, images: List[Image.Image]) -> AgentState:
    results = []

    for idx, image in enumerate(images[:2], start=1):
        try:
            img_b64 = encode_pil_image(image)

            response = vision_llm.invoke([
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"""
User request:
{state['user_query']}

Explain only the image details relevant for compliance review.
Focus on:
- key clauses
- highlighted obligations
- warning statements
- tables
- regulatory references
- risky or ambiguous language
"""
                        },
                        {
                            "type": "image_url",
                            "image_url": f"data:image/png;base64,{img_b64}"
                        }
                    ]
                }
            ])

            results.append(response.content)
        except Exception as e:
            pass

    state["vision_results"] = results
    return state


def reasoning_agent(state: AgentState) -> AgentState:
    retrieved_text = format_retrieved_chunks(
        state["retrieved_chunks"],
        state["retrieved_metadata"]
    )

    vision_text = "\n\n".join(
        [f"[Vision {i+1}]\n{item}" for i, item in enumerate(state["vision_results"])]
    ) if state["vision_results"] else "No vision results."

    csv_context = ""
    if state.get("document_type") == "csv":
        csv_context = f"""
Additional CSV review focus:
- identify missing mandatory fields
- identify suspicious blank values
- detect inconsistent value formatting
- identify duplicate or repeated records if visible in evidence
- flag unusual numeric or date patterns

CSV summary:
{json.dumps(state.get("csv_summary", {}), indent=2)}
"""

    prompt = f"""
You are a compliance risk analyst.

Document:
{state['document_name']}

Document type:
{state['document_type']}

User request:
{state['user_query']}

Plan:
{json.dumps(state['plan'], indent=2)}

Retrieved evidence:
{retrieved_text}

Vision evidence:
{vision_text}

{csv_context}

Your job:
- identify compliance risks
- highlight key clauses
- point out ambiguous or potentially non-compliant language
- summarize important findings
- suggest what a reviewer should inspect further

Rules:
- use only provided evidence
- do not invent regulations or clauses
- if evidence is insufficient, say so clearly
- produce a strong draft review
- structure the draft with:
1. Executive Summary
2. Key Clauses
3. Risk Flags
4. Potential Gaps / Ambiguities
5. Reviewer Recommendations
"""
    response = reasoning_llm.invoke(prompt)
    state["reasoning_output"] = response.content
    return state

def critic_agent(state: AgentState) -> AgentState:

    prompt = f"""
You are a strict critic for a rule-based compliance review system.

User request:
{state['user_query']}

Rules:
{json.dumps(state['active_rules'], indent=2)}

Rule findings:
{json.dumps(state['rule_findings'], indent=2)}

Draft review:
{state['reasoning_output']}

Return ONLY valid JSON:
{{
  "needs_revision": true,
  "issues": [],
  "missing_points": [],
  "unsupported_claims": [],
  "suggested_fixes": []
}}

Check:
- whether all active rules were addressed
- whether any matched finding is unsupported
- whether severity presentation is misleading
- whether important recommendations are missing
"""
    response = critic_llm.invoke(prompt)

    state["critic_feedback"] = safe_json_loads(
        response.content,
        fallback={
            "needs_revision": False,
            "issues": [],
            "missing_points": [],
            "unsupported_claims": [],
            "suggested_fixes": []
        }
    )

    return state

def rule_evaluation_agent(state: AgentState) -> AgentState:

    retrieved_text = format_retrieved_chunks(
        state["retrieved_chunks"],
        state["retrieved_metadata"]
    )

    findings = []

    for rule in state["active_rules"]:
        prompt = f"""
        You are a compliance rule evaluator.

        User request:
        {state['user_query']}

        Document:
        {state['document_name']}

        Rule:
        {json.dumps(rule, indent=2)}

        Retrieved evidence:
        {retrieved_text}

        Vision evidence:
        {json.dumps(state['vision_results'], indent=2)}

        Return ONLY valid JSON:
        {{
        "rule_id": "{rule.get('id', '')}",
        "rule_name": "{rule.get('name', '')}",
        "matched": true,
        "severity": "{rule.get('severity', 'Medium')}",
        "finding": "clear explanation",
        "evidence_summary": "relevant supporting evidence",
        "risk_reason": "why this matters",
        "reviewer_recommendation": "what reviewer should do next"
        }}

        Rules:
        - use only the provided evidence
        - if the rule is not supported by evidence, set matched=false
        - do not invent clauses
        """
        response = reasoning_llm.invoke(prompt)

        finding = safe_json_loads(
            response.content,
            fallback={
                "rule_id": rule.get("id", ""),
                "rule_name": rule.get("name", ""),
                "matched": False,
                "severity": rule.get("severity", "Medium"),
                "finding": "Unable to evaluate confidently from available evidence.",
                "evidence_summary": "",
                "risk_reason": "",
                "reviewer_recommendation": rule.get("reviewer_recommendation", "")
            }
        )
        findings.append(finding)

    existing = state.get("rule_findings", [])
    state["rule_findings"] = existing + findings
    return state

def revision_reasoning_agent(state: AgentState) -> AgentState:

    prompt = f"""
        Revise the compliance review draft using critic feedback.

        User request:
        {state['user_query']}

        Original draft:
        {state['reasoning_output']}

        Critic feedback:
        {json.dumps(state['critic_feedback'], indent=2)}

        Retrieved chunks:
        {json.dumps(state['retrieved_chunks'], indent=2)}

        Vision results:
        {json.dumps(state['vision_results'], indent=2)}

        Rules:
        - fix only grounded issues
        - do not invent any regulation or policy text
        - preserve structure
        - if evidence is not enough, explicitly say that
        """

    response = reasoning_llm.invoke(prompt)
    state["reasoning_output"] = response.content

    return state


def final_answer_agent(state: AgentState) -> AgentState:

    prompt = f"""
You are the final report generator for a compliance review assistant.

User request:
{state['user_query']}

Draft review:
{state['reasoning_output']}

Critic feedback:
{json.dumps(state['critic_feedback'], indent=2)}

Write a polished final compliance review report.

Requirements:
- make it clear and professional
- keep it precise in the draft
- include concise headings
- do not mention internal agents
- finish with a short reviewer action checklist
"""

    response = final_llm.invoke(prompt)
    state["final_answer"] = response.content

    return state


def run_multi_agent_pipeline(
    user_query: str,
    document_name: str,
    document_type: str,
    vectordb: Chroma,
    active_rules: List[Dict[str, Any]],
    extracted_images: Optional[List[Image.Image]] = None,
    csv_summary: Optional[Dict[str, Any]] = None,
    extracted_df: Optional[pd.DataFrame] = None
) -> AgentState:

    state: AgentState = {
        "user_query": user_query,
        "document_name": document_name,
        "document_type": document_type,
        "plan": {},
        "active_rules": active_rules,
        "retrieved_chunks": [],
        "retrieved_metadata": [],
        "vision_results": [],
        "rule_findings": [],
        "reasoning_output": "",
        "critic_feedback": {},
        "final_answer": "",
        "csv_summary": csv_summary
    }

    state = planner_agent(state)

    if state["plan"].get("needs_retrieval", True):
        state = retriever_agent(state, vectordb, k=5)

    if state["plan"].get("needs_vision", False) and extracted_images:
        state = vision_agent(state, extracted_images)

    if document_type == "csv" and extracted_df is not None:
        csv_rule_findings = evaluate_csv_structure_rules(extracted_df, active_rules)
        state["rule_findings"].extend(csv_rule_findings)

    state = rule_evaluation_agent(state)
    state = reasoning_agent(state)
    state = critic_agent(state)

    if state["critic_feedback"].get("needs_revision", False):
        state = revision_reasoning_agent(state)

    state = final_answer_agent(state)

    return state


st.subheader("Upload the document you need review")
uploaded_file = st.file_uploader(
    "Upload compliance document",
    type=["pdf", "docx", "txt","csv"]
)

st.subheader("Review Instruction")
default_query = (
    "Review this document for compliance risks, highlight key clauses, "
    "identify ambiguities, and provide reviewer recommendations."
)
user_query = st.text_area(
    "Review instruction",
    value=default_query,
    height=120
)

st.subheader("Compliance Rules")

def load_rules(file_path: str = "rules.json") -> List[Dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8") as f:
        return json.load(f)

try:
    default_rules = load_rules("rules.json")
except Exception:
    default_rules = [
        {
            "id": "R001",
            "name": "Data Retention Clause Missing",
            "category": "Data Governance",
            "severity": "High",
            "description": "Documents handling sensitive or regulated data should mention retention or disposal policy.",
            "required_keywords": ["retention", "retain", "storage period", "disposal"],
            "risk_if_missing": "Data lifecycle controls may be undefined.",
            "reviewer_recommendation": "Check whether a clear retention or disposal clause is included."
        },
        {
            "id": "R002",
            "name": "Ambiguous Approval Language",
            "category": "Process Control",
            "severity": "Medium",
            "description": "Approval responsibilities should be explicit and not vague.",
            "required_keywords": ["approved by", "authorization", "sign-off", "reviewed by"],
            "prohibited_patterns": ["as needed", "where appropriate", "if required"],
            "risk_if_missing": "Approval accountability may be unclear.",
            "reviewer_recommendation": "Ensure approval owner are explicitly defined."
        }
    ]

rules_mode = st.radio(
    "Rule source",
    ["Default Rules","Upload Rules PDF"],
    horizontal=True
)

active_rules = default_rules

if rules_mode == "Upload Rules JSON":
    rules_file = st.file_uploader(
        "Upload rules JSON",
        type=["json"],
        key="rules_upload"
    )
    if rules_file:
        try:
            active_rules = json.load(rules_file)
            st.success("Uploaded rules JSON loaded successfully.")
        except Exception as e:
            st.error(f"Invalid uploaded JSON: {e}")
            active_rules = default_rules

elif rules_mode == "Edit Rules JSON":
    rules_text = st.text_area(
        "Edit rules JSON",
        value=json.dumps(default_rules, indent=2),
        height=320
    )
    try:
        active_rules = json.loads(rules_text)
        st.success("Rules JSON is valid.")
    except Exception as e:
        st.error(f"Invalid JSON: {e}")
        active_rules = default_rules

elif rules_mode == "Upload Rules PDF":
    st.info(
        " Upload a PDF containing your compliance rules. "
        "The **DeepSeek agent** will read the PDF and automatically extract "
        "structured rules from it."
    )

    rules_pdf_file = st.file_uploader(
        "Upload rules PDF",
        type=["pdf"],
        key="rules_pdf_upload"
    )

    if rules_pdf_file:
        with st.spinner("🔍 Extracting text from rules PDF..."):
            try:
                rules_pdf_text = extract_text_from_rules_pdf(rules_pdf_file)
            except Exception as e:
                st.error(f"Failed to extract text from PDF: {e}")
                rules_pdf_text = ""

        if rules_pdf_text.strip():
            st.success(f" PDF text extracted ({len(rules_pdf_text)} characters).")

            with st.expander("Preview Extracted PDF Text", expanded=False):
                st.text(rules_pdf_text[:3000] + ("..." if len(rules_pdf_text) > 3000 else ""))

            with st.spinner(" DeepSeek agent is extracting compliance rules from PDF..."):
                extracted_rules = deepseek_rules_extraction_agent(rules_pdf_text)

            if extracted_rules:
                active_rules = extracted_rules
                st.success(f" DeepSeek extracted **{len(active_rules)} rules** from the PDF.")
            else:
                st.warning(
                    " DeepSeek could not extract structured rules. "
                    "Falling back to default rules."
                )
                active_rules = default_rules

            with st.expander(" Review / Edit Extracted Rules JSON", expanded=False):
                edited_rules_text = st.text_area(
                    "Edit extracted rules if needed",
                    value=json.dumps(active_rules, indent=2),
                    height=400,
                    key="extracted_rules_edit"
                )
                if st.button("Apply Edited Rules", key="apply_edited"):
                    try:
                        active_rules = json.loads(edited_rules_text)
                        st.success("Edited rules applied.")
                    except Exception as e:
                        st.error(f"Invalid JSON: {e}")
        else:
            st.error("Could not extract any text from the uploaded rules PDF.")


with st.expander("Preview Active Rules", expanded=False):
    st.json(active_rules)


run_clicked = st.button("Run Compliance Review", type="primary")

if run_clicked:
    st.session_state.agent_logs = []

    if not uploaded_file:
        st.error("Please upload a document.")
        st.stop()

    if not user_query.strip():
        st.error("Please provide a review instruction.")
        st.stop()

    if not isinstance(active_rules, list) or len(active_rules) == 0:
        st.error("Active rules are empty or invalid.")
        st.stop()

    with st.spinner("Reviewing document..."):

        extracted_text, extracted_images, extracted_df = extract_content(uploaded_file)
        # extracted_text, extracted_images = extract_content(uploaded_file)

        if not extracted_text.strip():
            st.error("No readable content extracted from the uploaded file.")
            st.stop()
            
        if extracted_df is not None:
            csv_summary = csv_precheck(extracted_df)
            st.subheader("CSV Quality Checks")
            st.json(csv_summary)
        else:
            csv_summary = None

        if extracted_df is not None:
            chunks, metadatas = build_csv_chunks(extracted_df, uploaded_file.name)
        else:
            chunks = split_text(extracted_text)
            metadatas = [
                {"source": uploaded_file.name, "chunk_id": i + 1}
                for i in range(len(chunks))
            ]

        metadatas = [
            {"source": uploaded_file.name, "chunk_id": i + 1}
            for i in range(len(chunks))
        ]
        vectordb = create_vector_db(chunks, metadatas)

        document_type = None
        if uploaded_file.name.lower().endswith(".pdf"):
            document_type = "pdf"
        elif uploaded_file.name.lower().endswith(".docx"):
            document_type = "docx"
        elif uploaded_file.name.lower().endswith(".txt"):
            document_type = "txt"
        elif uploaded_file.name.lower().endswith(".csv"):
            document_type = "csv"
        else:
            document_type = "unknown"

        state = run_multi_agent_pipeline(
            user_query=user_query.strip(),
            document_name=uploaded_file.name,
            document_type=document_type,
            vectordb=vectordb,
            active_rules=active_rules,
            extracted_images=extracted_images,
            csv_summary=csv_summary if extracted_df is not None else None,
            extracted_df=extracted_df if extracted_df is not None else None
        )


    st.success("Compliance review completed")


    st.subheader("Final Review Report")
    st.write(state["final_answer"])


    with st.expander("Planner Output"):
        st.json(state["plan"])

    with st.expander("Rule Findings", expanded=True):
        if state.get("rule_findings"):
            for idx, finding in enumerate(state["rule_findings"], start=1):
                st.markdown(
                    f"### Rule {idx}: {finding.get('rule_name', 'Unknown Rule')}"
                )
                st.write(f"**Rule ID:** {finding.get('rule_id', 'N/A')}")
                st.write(f"**Matched:** {finding.get('matched', False)}")
                st.write(f"**Severity:** {finding.get('severity', 'N/A')}")
                st.write(f"**Finding:** {finding.get('finding', '')}")
                st.write(f"**Evidence Summary:** {finding.get('evidence_summary', '')}")
                st.write(f"**Risk Reason:** {finding.get('risk_reason', '')}")
                st.write(
                    f"**Reviewer Recommendation:** "
                    f"{finding.get('reviewer_recommendation', '')}"
                )
                st.divider()
        else:
            st.info("No rule findings available.")

    with st.expander("Critic Feedback"):
        st.json(state["critic_feedback"])

    with st.expander("Retrieved Evidence"):
        if state.get("retrieved_chunks"):
            for i, chunk in enumerate(state["retrieved_chunks"], start=1):
                st.markdown(f"**Chunk {i}**")
                st.write(chunk[:1500])
        else:
            st.info("No retrieved evidence available.")

    if state.get("vision_results"):
        with st.expander("Vision Insights"):
            for i, item in enumerate(state["vision_results"], start=1):
                st.markdown(f"**Vision {i}**")
                st.write(item)
