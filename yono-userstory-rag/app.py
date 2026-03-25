import os 
import httpx 
import tiktoken 
import streamlit as st
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
import pytesseract
from PIL import Image
import io
import base64
import pdfplumber
import httpx 
import tiktoken 
tiktoken_cache_dir = "./token"
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir 

client = httpx.Client(verify=False) 
st.set_page_config(page_title="YONO AI Assistant", layout="wide")
st.title("📘 YONO AI – User Story Assistant")

llm = ChatOpenAI( 
base_url="https://genailab.tcs.in", 
model="azure_ai/genailab-maas-DeepSeek-V3-0324", 
api_key="sk-ngPOpvwa-2hjZ0iQZMpMUw", 
http_client=client
) 

llm_image = ChatOpenAI(model="azure/genailab-maas-gpt-4o ",
                 base_url="https://genailab.tcs.in",
                 api_key="sk-ngPOpvwa-2hjZ0iQZMpMUw",
                 http_client=client,
                 temperature= 0.7
)


def extract_image_with_llm(image):
    import io
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode()

    response = llm_image.invoke([
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "Extract and explain all information from this image in detail. Include flows, labels, steps, and relationships."},
                {"type": "image_url", "image_url": f"data:image/png;base64,{img_str}"}
            ]
        }
    ])
# Extract content from DOCX
def extract_docx(file):
    doc = Document(file)
    content = []

    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(row_text))

    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            image_data = doc.part.rels[rel]._target.blob
            image = Image.open(io.BytesIO(image_data))
            image_text = extract_image_with_llm(image)
            content.append("[Image Understanding]")
            content.append(image_text)
            content.append(text)

    return "\n".join(content)

# Extract content from PDF
def extract_pdf(file):
    content = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                content.append(text)
    return "\n".join(content)

# Split text
def split_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_text(text)

# Create vector DB
def create_vector_db(chunks):
    embeddings = OpenAIEmbeddings( 
base_url="https://genailab.tcs.in", 
model="azure/genailab-maas-text-embedding-3-large", 
api_key="sk-ngPOpvwa-2hjZ0iQZMpMUw", 
http_client=client) 
    vectordb = Chroma.from_texts(chunks, embedding=embeddings, persist_directory="./db")
    return vectordb

# Create QA chain
def create_qa(vectordb):
    retriever = vectordb.as_retriever(search_kwargs={"k": 4})
    qa = RetrievalQA.from_chain_type(
        llm = llm,
        retriever=retriever,
        return_source_documents=True
    )
    return qa

uploaded_file = st.file_uploader("Upload User Story DOCX or PDF", type=["docx", "pdf"])

if uploaded_file:
    with st.spinner("Processing document..."):
        if uploaded_file.name.lower().endswith(".pdf"):
            text = extract_pdf(uploaded_file)
        else:
            text = extract_docx(uploaded_file)
        chunks = split_text(text)
        vectordb = create_vector_db(chunks)
        qa = create_qa(vectordb)

    st.success("Document processed successfully!")

    query = st.text_input("Ask your question about user stories")

    if query:
        result = qa({"query": query})
        st.subheader("Answer")
        st.write(result["result"])

        st.subheader("Sources")
        for doc in result["source_documents"]:
            st.write(doc.page_content[:300])
