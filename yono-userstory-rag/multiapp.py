import os 
import httpx 
import tiktoken 
import streamlit as st
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_classic.chains import RetrievalQA
import pytesseract
from PIL import Image
import io
import base64
import pdfplumber
import httpx 
import tiktoken 
import re
import json
import json5
from openai import OpenAI
from typing import TypedDict, Dict, Any, List, Optional
import time


BASE_URL = "https://genailab.tcs.in"
API_KEY = "sk-ngPOpvwa-2hjZ0iQZMpMUw"


# from streamlit_mic_recorder import mic_recorder

tiktoken_cache_dir = "./token"
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir 

client = httpx.Client(verify=False) 
st.set_page_config(page_title="YONO AI Assistant", layout="wide")
st.title("📘 YONO AI – User Story Assistant")

llm = ChatOpenAI( 
base_url="https://genailab.tcs.in", 
model="azure_ai/genailab-maas-DeepSeek-V3-0324", 
api_key="sk-ngPOpvwa-2hjZ0iQZMpMUw", 
http_client=client
) 


llm_image = ChatOpenAI(model="azure/genailab-maas-gpt-4o ",
                 base_url="https://genailab.tcs.in",
                 api_key="sk-ngPOpvwa-2hjZ0iQZMpMUw",
                 http_client=client,
                 temperature= 0.7
)

planner_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure/genailab-maas-gpt-4o",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

reasoning_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure_ai/genailab-maas-DeepSeek-V3-0324",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

critic_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure_ai/genailab-maas-DeepSeek-V3-0324",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

final_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure/genailab-maas-gpt-4o-mini",
    api_key=API_KEY,
    http_client=client,
    temperature=0.2
)

vision_llm = ChatOpenAI(
    base_url=BASE_URL,
    model="azure/genailab-maas-gpt-4o",
    api_key=API_KEY,
    http_client=client,
    temperature=0
)

embedding_model = OpenAIEmbeddings(
    base_url=BASE_URL,
    model="azure/genailab-maas-text-embedding-3-large",
    api_key=API_KEY,
    http_client=client
)


class AgentState(TypedDict):
    user_query: str
    plan: Dict[str, Any]
    retrieved_chunks: List[str]
    retrieved_metadata: List[Dict[str, Any]]
    vision_results: List[str]
    reasoning_output: str
    critic_feedback: Dict[str, Any]
    final_answer: str

def safe_json_loads(text: str, fallback: Optional[dict] = None) -> dict:
    try:
        return json.loads(text)
    except Exception:
        return fallback or {}

def mask_sensitive_data(text: str):
    patterns = {
        "EMAIL": r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+',
        "PHONE": r'\b\d{10}\b',
        "AADHAAR": r'\b\d{4}\s?\d{4}\s?\d{4}\b',
        "PAN": r'\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b',
        "CREDIT_CARD": r'\b\d{4}[- ]?\d{4}[- ]?\d{4}[- ]?\d{4}\b',
        "ACCOUNT_NUMBER": r'\b\d{9,18}\b'
    }

    detected = []

    for label, pattern in patterns.items():
        matches = re.findall(pattern, text)
        if matches:
            detected.extend([(label, m) for m in matches])
            text = re.sub(pattern, f"[REDACTED_{label}]", text)

    return text, detected

def extract_image_with_llm(image):
    import io
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode()

    response = llm_image.invoke([
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "Extract and explain all information from this image in detail. Include flows, labels, steps, and relationships."},
                {"type": "image_url", "image_url": f"data:image/png;base64,{img_str}"}
            ]
        }
    ])
# Extract content from DOCX
def extract_docx(file):
    doc = Document(file)
    content = []

    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(row_text))

    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            image_data = doc.part.rels[rel]._target.blob
            image = Image.open(io.BytesIO(image_data))
            image_text = extract_image_with_llm(image)
            content.append("[Image Understanding]")
            content.append(image_text)
            content.append(extracted_text)

    return "\n".join(content)

# Extract content from PDF
def extract_pdf(file):
    content = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                content.append(text)
    return "\n".join(content)

# Split text
def split_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_text(text)

# Create vector DB
# def create_vector_db(chunks):
#     embeddings = OpenAIEmbeddings( 
#     base_url="https://genailab.tcs.in", 
#     model="azure/genailab-maas-text-embedding-3-large", 
#     api_key="sk-ngPOpvwa-2hjZ0iQZMpMUw", 
#     http_client=client
#     ) 
#     vectordb = Chroma.from_texts(chunks, embedding=embeddings, persist_directory="./db")
#     return vectordb
def create_vector_db(chunks: List[str], metadatas: Optional[List[dict]] = None):
    # persist_dir = os.path.join(tempfile.gettempdir(), f"yono_chroma_{uuid.uuid4().hex}")
    return Chroma.from_texts(
        texts=chunks,
        embedding=embedding_model,
        metadatas=metadatas or [{} for _ in chunks],
        # persist_directory=persist_dir
    )

# Create QA chain
def create_qa(vectordb):
    retriever = vectordb.as_retriever(search_kwargs={"k": 4})
    qa = RetrievalQA.from_chain_type(
        llm = llm,
        retriever=retriever,
        return_source_documents=True
    )
    return qa

def extract_content(uploaded_file) -> tuple[str, List[Image.Image]]:
    if uploaded_file.name.lower().endswith(".pdf"):
        return extract_pdf(uploaded_file),[]
    return extract_docx(uploaded_file),[]



def planner_agent(state: AgentState) -> AgentState:
    # Implement planning logic here
    prompt = f"""You are a planner agent for an enterprise document assistant.
    
    User query: {state['user_query']}
    Return ONLY valid JSON:
    {{
    "needs_retrieval": true,
    "needs_vision": false,
    "task_breakdown": ["task 1", "task 2","..."],
    "answer_style": "concise|detailed|stepwise"
    }}

    Set needs_vision=true if the user asks about:
    figure, image, screenshot, UI, diagram, flow, journey map
    Plan:"""
    
    response = planner_llm.invoke(prompt)
    
    state["plan"] = safe_json_loads(response.content, fallback={
        "needs_retrieval": True,
        "needs_vision": False,
        "task_breakdown": ["answer the user question from the document"],
        "answer_style": "detailed"
    })
    return state

def retriever_agent(state: AgentState, vectordb: Chroma, k :int = 5) -> AgentState:
    # Implement retrieval logic here
    retriever = vectordb.as_retriever(search_kwargs={"k": k})
    docs = retriever.invoke(state["user_query"])
    state["retrieved_chunks"] = [doc.page_content for doc in docs]
    state["retrieved_metadata"] = [doc.metadata for doc in docs]
    return state

def vision_agent(state: AgentState, extracted_images: Optional[List[Image.Image]]) -> AgentState:
    # Implement vision logic here
    return state

def reasoning_agent(state: AgentState) -> AgentState:
    # Implement reasoning logic here
    retrieved_text = "\n\n".join(
        [f"[Chunk {i+1}]\n{chunk}" for i, chunk in enumerate(state["retrieved_chunks"])]
    )
    state["reasoning_output"] = reasoning_llm.invoke(retrieved_text).content
    return state

def critic_agent(state: AgentState) -> AgentState:
    # Implement critic logic here
    prompt = f"""
        You are a strict critic agent.

        User query:
        {state['user_query']}

        Plan:
        {json.dumps(state['plan'], indent=2)}

        Draft answer:
        {state['reasoning_output']}

        Retrieved chunks:
        {json.dumps(state['retrieved_chunks'], indent=2)}

        Vision results:
        {json.dumps(state['vision_results'], indent=2)}

        Return ONLY valid JSON:
        {{
        "needs_revision": true,
        "issues": ["issue"],
        "missing_points": ["missing point"],
        "suggested_fixes": ["fix"]
        }}

        Mark needs_revision=true if:
        - any planned task is not covered
        - unsupported claims exist
        - answer is vague
        - image-related questions are not handled properly
        """
    response = critic_llm.invoke(prompt)
    state["critic_feedback"] = safe_json_loads(response.content, fallback={
        "needs_revision": False,
        "issues": [],
        "missing_points": [],
        "suggested_fixes": []
    })
    print(state["critic_feedback"])
    return state

def revision_reasoning_agent(state: AgentState) -> AgentState:
    # Implement revision reasoning logic here
    prompt = f"""
        Revise the draft answer using the critic feedback.

        User query:
        {state['user_query']}

        Original draft:
        {state['reasoning_output']}

        Critic feedback:
        {json.dumps(state['critic_feedback'], indent=2)}

        Retrieved chunks:
        {json.dumps(state['retrieved_chunks'], indent=2)}

        Vision results:
        {json.dumps(state['vision_results'], indent=2)}

        Rules:
        - fix all grounded issues
        - do not invent facts
        - mention missing evidence if needed
        """
    response = reasoning_llm.invoke(prompt)
    state["reasoning_output"] = response.content
    print("Final answer agent completed")
    return state

def final_answer_agent(state: AgentState) -> AgentState:
    # Implement final answer logic here
    prompt = f"""
        You are the final response agent.

        User query:
        {state['user_query']}

        Draft answer:
        {state['reasoning_output']}

        Critic feedback:
        {json.dumps(state['critic_feedback'], indent=2)}

        Write the final answer for the user.
        Rules:
        - be polished
        - be structured
        - stay grounded in evidence
        - do not mention internal agent workflow
    """
    response = final_llm.invoke(prompt)
    state["final_answer"] = response.content
    print("Multi-agent pipeline completed")
    return state

def run_multi_agent_pipeline(
    user_query: str,
    vectordb: Chroma,
    extracted_images: Optional[List[Image.Image]] = None
) -> AgentState:
    state: AgentState = {
        "user_query": user_query,
        "plan": {},
        "retrieved_chunks": [],
        "retrieved_metadata": [],
        "vision_results": [],
        "reasoning_output": "",
        "critic_feedback": {},
        "final_answer": ""
    }

    state = planner_agent(state)
    print(state["plan"])
    if state["plan"].get("needs_retrieval", True):
        state = retriever_agent(state, vectordb)

    if state["plan"].get("needs_vision", False) and extracted_images:
        state = vision_agent(state, extracted_images)

    state = reasoning_agent(state)
    state = critic_agent(state)

    if state["critic_feedback"].get("needs_revision", False):
        state = revision_reasoning_agent(state)
    state["critic_feedback"] = {
    "needs_revision": False,
    "issues": [],
    "missing_points": [],
    "suggested_fixes": []
}

    state = final_answer_agent(state)
    return state



uploaded_file = st.file_uploader("Upload User Story DOCX or PDF", type=["docx", "pdf"])

if uploaded_file:
    with st.spinner("Processing document..."):
        extracted_text, extracted_images = extract_content(uploaded_file)
        clean_text, detected_items = mask_sensitive_data(extracted_text)

        if detected_items:
            st.warning("Sensitive information found in document and masked.")

        chunks = split_text(clean_text)
        metadatas = [{"source": uploaded_file.name, "chunk_id": i + 1} for i in range(len(chunks))]
        vectordb = create_vector_db(chunks, metadatas)

    st.success("Document processed successfully.")

    query = st.text_input("Ask your question about user stories")

    if query:
        clean_query, detected = mask_sensitive_data(query)
        if detected:
            st.warning("Sensitive information detected in your query and has been masked.")

        with st.spinner("Agents are collaborating..."):
            state = run_multi_agent_pipeline(
                user_query=clean_query,
                vectordb=vectordb,
                extracted_images=extracted_images
            )

        answer_text = state["final_answer"]

        st.subheader("Answer")
        st.write(answer_text)

        # try:
        #     audio_path = text_to_speech_file(answer_text)
        #     with open(audio_path, "rb") as f:
        #         st.audio(f.read(), format="audio/mp3")
        # except Exception as e:
        #     st.error(f"Voice output failed: {type(e).__name__}: {e}")

        with st.expander("Planner Output"):
            st.json(state["plan"])

        with st.expander("Critic Feedback"):
            st.json(state["critic_feedback"])

        with st.expander("Retrieved Chunks"):
            for i, chunk in enumerate(state["retrieved_chunks"], start=1):
                st.markdown(f"**Chunk {i}**")
                st.write(chunk[:1200])

        if state["vision_results"]:
            with st.expander("Vision Insights"):
                for i, item in enumerate(state["vision_results"], start=1):
                    st.markdown(f"**Vision {i}**")
                    st.write(item)
                    
                    
                    
                    
# if uploaded_file:
#     with st.spinner("Processing document..."):
#         if uploaded_file.name.lower().endswith(".pdf"):
#             extracted_text = extract_pdf(uploaded_file)
#         else:
#             extracted_text = extract_docx(uploaded_file)
#         clean_text, detected_items = mask_sensitive_data(extracted_text)

#         if detected_items:
#             print("Sensitive data found and masked:", detected_items)

#         chunks = split_text(clean_text)

#         vectordb = create_vector_db(chunks)
#         qa = create_qa(vectordb)

#     st.success("Document processed successfully!")

#     query = st.text_input("Ask your question about user stories")

#     if query:
#         clean_query, detected = mask_sensitive_data(query)

#         if detected:
#             st.warning("⚠️ Sensitive information detected in your query and has been masked.")

#         result = qa.invoke({"query": clean_query})

#         answer_text = result["result"]
#         st.subheader("Answer")
#         st.write(answer_text)
#         print(result)
#         # st.subheader("Sources")
#         # for doc in result["source_documents"]:
#         #     st.write(doc.page_content[:300])
        
        
#         #text to speech
#         # try:
#         #     answer_audio = text_to_speech_bytes(answer_text, voice="alloy")
#         #     st.audio(answer_audio, format="audio/mp3")
#         # except Exception as e:
#         #     st.error(f"Voice output failed: {e}")

